from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from langgraph.graph import Graph, StateGraph
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage
import json
import asyncio
from typing_extensions import TypedDict
import dotenv
import os
from PyPDF2 import PdfReader
import io
from docx import Document
from zipfile import ZipFile
import tempfile
import logging


# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Pydantic Models
class SalaryEstimate(BaseModel):
    min: float = Field(default=0.0)
    max: float = Field(default=0.0)
    currency: str = Field(default="USD")

class ResumeResponse(BaseModel):
    name: str = Field(default="Unknown")
    title: str = Field(default="Professional")
    location: str = Field(default="Not specified")
    yearsOfExperience: int = Field(default=0)
    skills: List[str] = Field(default_factory=list)
    workPreference: str = Field(default="Remote")
    education: str = Field(default="Not specified")
    pastCompanies: List[str] = Field(default_factory=list)
    summary: str = Field(default="No summary available")
    overallScore: float = Field(default=0.0)
    strengths: List[str] = Field(default_factory=list)
    roleRecommendations: List[str] = Field(default_factory=list)
    salaryEstimate: SalaryEstimate = Field(default_factory=lambda: SalaryEstimate())
    topSkills: List[str] = Field(default_factory=list)

class ErrorResponse(BaseModel):
    error: str
    message: str

# LangGraph State
class ResumeParsingState(TypedDict):
    resume_text: str
    parsed_data: Optional[Dict[str, Any]]
    extraction_complete: bool
    error: Optional[str]

class ResumeParserGraph:
    def __init__(self, gemini_api_key: str):
        if not gemini_api_key:
            raise ValueError("GEMINI_KEY environment variable is not set")
        self.llm = ChatGoogleGenerativeAI(model='models/gemini-2.0-flash-lite', api_key=gemini_api_key)
        self.graph = self._create_graph()
    
    def _create_graph(self) -> StateGraph:
        # Create the state graph
        workflow = StateGraph(ResumeParsingState)
        
        # Add nodes
        workflow.add_node("extract_resume_data", self._extract_resume_data)
        workflow.add_node("validate_and_format", self._validate_and_format)
        workflow.add_node("handle_error", self._handle_error)
        
        # Set entry point
        workflow.set_entry_point("extract_resume_data")
        
        # Add edges
        workflow.add_conditional_edges(
            "extract_resume_data",
            self._should_validate,
            {
                "validate": "validate_and_format",
                "error": "handle_error"
            }
        )
        
        workflow.add_edge("validate_and_format", "__end__")
        workflow.add_edge("handle_error", "__end__")
        
        return workflow.compile()
    
    def _extract_resume_data(self, state: ResumeParsingState) -> ResumeParsingState:
        """Extract structured data from resume text using LLM"""
        try:
            # Clean and normalize the text
            resume_text = state['resume_text'].strip()
            if not resume_text:
                logger.error("Empty resume text received")
                state["error"] = "Empty resume text"
                state["extraction_complete"] = False
                return state

            logger.info(f"Processing resume text of length: {len(resume_text)}")
            
            prompt = f"""
            Parse the following resume text and extract information in JSON format with these exact fields:
            - name: Full name of the candidate
            - title: Professional title or current role
            - location: Current location (city, country format)
            - yearsOfExperience: Total years of professional experience (integer)
            - skills: Array of technical and professional skills
            - workPreference: "Remote", "Onsite", or "Hybrid" (infer from resume or default to "Remote")
            - education: Highest education qualification 
            - pastCompanies: Array of previous company names
            - summary: Brief professional summary (2-3 sentences)
            - overallScore: Score from 0–100 evaluating the resume's strength
            - strengths: Array of strengths or positive attributes
            - roleRecommendations: Array of roles the candidate may be fit for
            - salaryEstimate: JSON with min (number), max (number), and currency (string) keys
            - topSkills: Array of top 3–5 most relevant skills

            Resume text:
            {resume_text}

            Return only valid JSON without any additional text or markdown formatting.
            """
            
            response = self.llm.invoke([HumanMessage(content=prompt)])
            logger.info(f"LLM Response received: {response.content[:200]}...")  # Log first 200 chars
            
            # Try to parse the JSON response
            try:
                # Clean the response content
                content = response.content.strip()
                # Remove any markdown code block markers
                if content.startswith("```json"):
                    content = content[7:-3]
                elif content.startswith("```"):
                    content = content[3:-3]
                
                # Try to parse the cleaned content
                parsed_data = json.loads(content)
                logger.info(f"Successfully parsed JSON response")
                
                # Ensure all required fields are present with defaults
                formatted_data = {
                    "name": parsed_data.get("name", "Unknown"),
                    "title": parsed_data.get("title", "Professional"),
                    "location": parsed_data.get("location", "Not specified"),
                    "yearsOfExperience": int(parsed_data.get("yearsOfExperience", 0)),
                    "skills": parsed_data.get("skills", []) if isinstance(parsed_data.get("skills"), list) else [],
                    "workPreference": parsed_data.get("workPreference", "Remote"),
                    "education": parsed_data.get("education", "Not specified"),
                    "pastCompanies": parsed_data.get("pastCompanies", []) if isinstance(parsed_data.get("pastCompanies"), list) else [],
                    "summary": parsed_data.get("summary", "No summary available"),
                    "overallScore": float(parsed_data.get("overallScore", 0.0)),
                    "strengths": parsed_data.get("strengths", []) if isinstance(parsed_data.get("strengths"), list) else [],
                    "roleRecommendations": parsed_data.get("roleRecommendations", []) if isinstance(parsed_data.get("roleRecommendations"), list) else [],
                    "salaryEstimate": {
                        "min": float(parsed_data.get("salaryEstimate", {}).get("min", 0.0)),
                        "max": float(parsed_data.get("salaryEstimate", {}).get("max", 0.0)),
                        "currency": parsed_data.get("salaryEstimate", {}).get("currency", "USD")
                    },
                    "topSkills": parsed_data.get("topSkills", []) if isinstance(parsed_data.get("topSkills"), list) else []
                }
                
                state["parsed_data"] = formatted_data
                state["extraction_complete"] = True
                logger.info("Successfully extracted and formatted resume data")
                
            except json.JSONDecodeError as e:
                logger.error(f"JSON parsing error: {str(e)}")
                logger.error(f"Failed to parse content: {content}")
                state["error"] = f"Failed to parse LLM response as JSON: {str(e)}"
                state["extraction_complete"] = False
            
        except Exception as e:
            logger.error(f"Error during extraction: {str(e)}")
            state["error"] = f"Error during extraction: {str(e)}"
            state["extraction_complete"] = False
        
        return state
    
    def _validate_and_format(self, state: ResumeParsingState) -> ResumeParsingState:
        """Validate and format the extracted data"""
        try:
            data = state["parsed_data"]
            
            # Ensure all required fields are present with defaults
            formatted_data = {
                "name": data.get("name", "Unknown"),
                "title": data.get("title", "Professional"),
                "location": data.get("location", "Not specified"),
                "yearsOfExperience": int(data.get("yearsOfExperience", 0)),
                "skills": data.get("skills", []) if isinstance(data.get("skills"), list) else [],
                "workPreference": data.get("workPreference", "Remote"),
                "education": data.get("education", "Not specified"),
                "pastCompanies": data.get("pastCompanies", []) if isinstance(data.get("pastCompanies"), list) else [],
                "summary": data.get("summary", "No summary available"),
                "overallScore": float(data.get("overallScore", 0.0)),
                "strengths": data.get("strengths", []) if isinstance(data.get("strengths"), list) else [],
                "roleRecommendations": data.get("roleRecommendations", []) if isinstance(data.get("roleRecommendations"), list) else [],
                "salaryEstimate": {
                    "min": float(data.get("salaryEstimate", {}).get("min", 0)),
                    "max": float(data.get("salaryEstimate", {}).get("max", 0)),
                    "currency": data.get("salaryEstimate", {}).get("currency", "USD")
                },
                "topSkills": data.get("topSkills", []) if isinstance(data.get("topSkills"), list) else []
            }

            
            # Validate workPreference
            if formatted_data["workPreference"] not in ["Remote", "Onsite", "Hybrid"]:
                formatted_data["workPreference"] = "Remote"
            
            state["parsed_data"] = formatted_data
            
        except Exception as e:
            state["error"] = f"Error during validation: {str(e)}"
            state["extraction_complete"] = False
        
        return state
    
    def _handle_error(self, state: ResumeParsingState) -> ResumeParsingState:
        """Handle errors in the parsing process"""
        # Create a default response structure for errors
        state["parsed_data"] = {
            "name": "N/A",
            "title": "N/A",
            "location": "N/A",
            "yearsOfExperience": 0,
            "skills": [],
            "workPreference": "Remote",
            "education": "N/A",
            "pastCompanies": [],
            "summary": f"Error parsing resume: {state.get('error', 'Unknown error')}",
            "overallScore": 0.0,
            "strengths": [],
            "roleRecommendations": [],
            "salaryEstimate": {"min": 0, "max": 0, "currency": "USD"},
            "topSkills": []
        }

        return state
    
    def _should_validate(self, state: ResumeParsingState) -> str:
        """Determine next step based on extraction success"""
        if state["extraction_complete"] and not state.get("error"):
            return "validate"
        else:
            return "error"
    
    async def parse_resume(self, resume_text: str) -> Dict[str, Any]:
        """Parse resume text and return structured data"""
        try:
            logger.info(f"Starting resume parsing for text of length: {len(resume_text)}")
            initial_state = ResumeParsingState(
                resume_text=resume_text,
                parsed_data=None,
                extraction_complete=False,
                error=None
            )
            
            # Run the graph
            result = await asyncio.get_event_loop().run_in_executor(
                None, self.graph.invoke, initial_state
            )
            
            if result.get("error"):
                logger.error(f"Error in resume parsing: {result['error']}")
                return None
            
            if not result.get("parsed_data"):
                logger.error("No parsed data returned from resume parsing")
                return None
            
            logger.info("Successfully parsed resume")
            return result["parsed_data"]
            
        except Exception as e:
            logger.error(f"Error in parse_resume: {str(e)}")
            return None

# FastAPI Application
app = FastAPI(
    title="Resume Parser API",
    description="API to parse resumes using LangGraph and return structured data",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Initialize the resume parser
dotenv.load_dotenv()
GEMINI_KEY = os.getenv("GEMINI_KEY")
if not GEMINI_KEY:
    raise ValueError("GEMINI_KEY environment variable is not set")
resume_parser = ResumeParserGraph(GEMINI_KEY)

@app.post("/parse-resume", response_model=List[ResumeResponse])
async def parse_resume_endpoint(file: UploadFile = File(...)):
    """
    Parse a resume file or ZIP archive of resumes and return structured data.
    
    Accepts: PDF, TXT, DOC, DOCX files or a ZIP containing them.
    Returns: List of structured resume data in JSON format.
    """
    try:
        logger.info(f"Received file: {file.filename} with content type: {file.content_type}")
        content = await file.read()
        resumes = []

        # If the file is a zip, extract supported files and process each
        if file.content_type == "application/zip" or file.filename.lower().endswith('.zip'):
            logger.info("Processing ZIP file")
            with tempfile.TemporaryDirectory() as tmpdir:
                zip_path = os.path.join(tmpdir, "resumes.zip")
                with open(zip_path, "wb") as f:
                    f.write(content)

                with ZipFile(zip_path, "r") as zip_ref:
                    file_list = zip_ref.namelist()
                    logger.info(f"Files in ZIP: {file_list}")
                    zip_ref.extractall(tmpdir)

                    for name in file_list:
                        file_path = os.path.join(tmpdir, name)
                        try:
                            logger.info(f"Processing file: {name}")
                            
                            if name.lower().endswith(".pdf"):
                                reader = PdfReader(file_path)
                                resume_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                                logger.info(f"Extracted {len(resume_text)} characters from PDF")
                            elif name.lower().endswith(".docx"):
                                doc = Document(file_path)
                                resume_text = "\n".join([para.text for para in doc.paragraphs])
                                logger.info(f"Extracted {len(resume_text)} characters from DOCX")
                            elif name.lower().endswith(".txt"):
                                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                                    resume_text = f.read()
                                logger.info(f"Extracted {len(resume_text)} characters from TXT")
                            else:
                                logger.warning(f"Skipping unsupported file type: {name}")
                                continue

                            if not resume_text.strip():
                                logger.warning(f"No text content extracted from {name}")
                                continue

                            logger.info(f"Parsing resume text from {name}")
                            parsed_data = await resume_parser.parse_resume(resume_text)
                            
                            if parsed_data and isinstance(parsed_data, dict):
                                logger.info(f"Successfully parsed {name}")
                                # Ensure all required fields are present with defaults
                                formatted_data = {
                                    "name": parsed_data.get("name", "Unknown"),
                                    "title": parsed_data.get("title", "Professional"),
                                    "location": parsed_data.get("location", "Not specified"),
                                    "yearsOfExperience": int(parsed_data.get("yearsOfExperience", 0)),
                                    "skills": parsed_data.get("skills", []) if isinstance(parsed_data.get("skills"), list) else [],
                                    "workPreference": parsed_data.get("workPreference", "Remote"),
                                    "education": parsed_data.get("education", "Not specified"),
                                    "pastCompanies": parsed_data.get("pastCompanies", []) if isinstance(parsed_data.get("pastCompanies"), list) else [],
                                    "summary": parsed_data.get("summary", "No summary available"),
                                    "overallScore": float(parsed_data.get("overallScore", 0.0)),
                                    "strengths": parsed_data.get("strengths", []) if isinstance(parsed_data.get("strengths"), list) else [],
                                    "roleRecommendations": parsed_data.get("roleRecommendations", []) if isinstance(parsed_data.get("roleRecommendations"), list) else [],
                                    "salaryEstimate": {
                                        "min": float(parsed_data.get("salaryEstimate", {}).get("min", 0.0)),
                                        "max": float(parsed_data.get("salaryEstimate", {}).get("max", 0.0)),
                                        "currency": parsed_data.get("salaryEstimate", {}).get("currency", "USD")
                                    },
                                    "topSkills": parsed_data.get("topSkills", []) if isinstance(parsed_data.get("topSkills"), list) else []
                                }
                                resumes.append(ResumeResponse(**formatted_data))
                            else:
                                logger.warning(f"No valid parsed data returned for {name}")
                        except Exception as e:
                            logger.error(f"Error processing {name}: {str(e)}")
                            continue

        else:
            # Process single file
            logger.info("Processing single file")
            try:
                if file.content_type == "text/plain":
                    resume_text = content.decode("utf-8")
                elif file.content_type == "application/pdf":
                    reader = PdfReader(io.BytesIO(content))
                    resume_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                elif file.content_type in [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/msword"
                ]:
                    doc = Document(io.BytesIO(content))
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    resume_text = content.decode("utf-8", errors="ignore")

                if not resume_text.strip():
                    raise HTTPException(status_code=400, detail="No text content could be extracted from the file")

                logger.info(f"Parsing resume text of length: {len(resume_text)}")
                parsed_data = await resume_parser.parse_resume(resume_text)
                
                if parsed_data and isinstance(parsed_data, dict):
                    logger.info("Successfully parsed single file")
                    # Ensure all required fields are present with defaults
                    formatted_data = {
                        "name": parsed_data.get("name", "Unknown"),
                        "title": parsed_data.get("title", "Professional"),
                        "location": parsed_data.get("location", "Not specified"),
                        "yearsOfExperience": int(parsed_data.get("yearsOfExperience", 0)),
                        "skills": parsed_data.get("skills", []) if isinstance(parsed_data.get("skills"), list) else [],
                        "workPreference": parsed_data.get("workPreference", "Remote"),
                        "education": parsed_data.get("education", "Not specified"),
                        "pastCompanies": parsed_data.get("pastCompanies", []) if isinstance(parsed_data.get("pastCompanies"), list) else [],
                        "summary": parsed_data.get("summary", "No summary available"),
                        "overallScore": float(parsed_data.get("overallScore", 0.0)),
                        "strengths": parsed_data.get("strengths", []) if isinstance(parsed_data.get("strengths"), list) else [],
                        "roleRecommendations": parsed_data.get("roleRecommendations", []) if isinstance(parsed_data.get("roleRecommendations"), list) else [],
                        "salaryEstimate": {
                            "min": float(parsed_data.get("salaryEstimate", {}).get("min", 0.0)),
                            "max": float(parsed_data.get("salaryEstimate", {}).get("max", 0.0)),
                            "currency": parsed_data.get("salaryEstimate", {}).get("currency", "USD")
                        },
                        "topSkills": parsed_data.get("topSkills", []) if isinstance(parsed_data.get("topSkills"), list) else []
                    }
                    resumes.append(ResumeResponse(**formatted_data))
                else:
                    logger.warning("No valid parsed data returned for single file")
            except Exception as e:
                logger.error(f"Error processing single file: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

        if not resumes:
            logger.error("No valid resumes found in the uploaded file(s)")
            raise HTTPException(status_code=400, detail="No valid resumes found in the uploaded file(s).")

        logger.info(f"Successfully parsed {len(resumes)} resumes")
        return resumes

    except Exception as e:
        logger.error(f"Error in parse_resume_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing resume(s): {str(e)}")

@app.post("/parse-resume-text", response_model=ResumeResponse)
async def parse_resume_text_endpoint(resume_text: str):
    """
    Parse resume text directly and return structured data
    
    Accepts: Raw resume text as string
    Returns: Structured resume data in JSON format
    """
    try:
        # Parse the resume using LangGraph
        parsed_data = await resume_parser.parse_resume(resume_text)
        
        # Create and return the response
        return ResumeResponse(**parsed_data)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "message": "Resume Parser API is running"}

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Resume Parser API",
        "version": "1.0.0",
        "endpoints": {
            "POST /parse-resume": "Upload resume file for parsing",
            "POST /parse-resume-text": "Send resume text for parsing",
            "GET /health": "Health check",
            "GET /docs": "API documentation"
        }
    }

# Example usage and testing
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, port=8000)